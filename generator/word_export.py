"""
Génère le document Word de synthèse DIP (8 sections).
Formaté pour être copié-collé en PDF/Word sans modification.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time


# ── Helpers de mise en forme ──────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def _bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    return p


def _table_2col(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Indicateur"
    hdr[1].text = "Valeur"
    for cell in hdr:
        run = cell.paragraphs[0].runs
        if run:
            run[0].bold = True
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value not in (None, "") else "N/D"
    doc.add_paragraph()


def _section_sep(doc: Document) -> None:
    doc.add_paragraph()


# ── Sections ──────────────────────────────────────────────────────────────────

def _section_zone(doc, zone_nom, dept_code, dept_nom, region, communes):
    _heading(doc, "1. Zone géographique", 1)
    _bullet(doc, f"Département : {dept_nom} ({dept_code})")
    _bullet(doc, f"Région : {region}")
    _bullet(doc, f"Nom de zone : {zone_nom}")
    noms = [f"{c['nom']} ({c['cp']})" for c in communes]
    _bullet(doc, f"Communes couvertes ({len(communes)}) : {', '.join(noms)}")
    _section_sep(doc)


def _section_socio_demo(doc, communes_data):
    _heading(doc, "2. Données socio-démographiques", 1)

    # Tableau par commune
    _heading(doc, "Données par commune (source : INSEE RP 2022)", 2)
    for r in communes_data:
        nom = r.get("commune", "")
        statut = r.get("_statut", "OK")
        _heading(doc, f"{nom}" + (" ⚠ " + statut if statut != "OK" else ""), 3)
        rows = [
            ("Population 2022", r.get("pop_2022")),
            ("Superficie (km²)", r.get("superficie_km2")),
            ("Densité (hab/km²) 2022", r.get("densite_2022")),
            ("Variation pop. 2016-2022 (%)", r.get("var_pop_2016_2022")),
            ("Dont solde naturel (%)", r.get("solde_naturel")),
            ("Dont solde migratoire (%)", r.get("solde_migratoire")),
            ("Nb ménages 2022", r.get("nb_menages_2022")),
            ("Naissances 2022", r.get("naissances_2022")),
            ("Décès 2022", r.get("deces_2022")),
            ("Nb logements 2022", r.get("nb_logements_2022")),
            ("Résidences principales (%)", r.get("part_res_principales_2022")),
            ("Résidences secondaires (%)", r.get("part_res_secondaires_2022")),
            ("Logements vacants (%)", r.get("part_logements_vacants_2022")),
            ("Ménages propriétaires (%)", r.get("part_proprietaires_2022")),
            ("Nb ménages fiscaux 2021", r.get("nb_menages_fiscaux_2021")),
            ("Ménages imposés 2021 (%)", r.get("part_menages_imposes_2021")),
            ("Médiane revenu disponible 2021 (€)", r.get("mediane_revenu_2021")),
            ("Taux de pauvreté 2021 (%)", r.get("taux_pauvrete_2021")),
            ("Emploi total 2022", r.get("emploi_total_2022")),
            ("Part emploi salarié 2022 (%)", r.get("part_emploi_salarie_2022")),
            ("Variation emploi 2016-2022 (%)", r.get("var_emploi_2016_2022")),
            ("Taux d'activité 15-64 ans 2022 (%)", r.get("taux_activite_15_64_2022")),
            ("Taux de chômage 15-64 ans 2022 (%)", r.get("taux_chomage_15_64_2022")),
            ("Nb établissements actifs 2023", r.get("nb_etab_actifs_2023")),
            ("Agriculture (%)", r.get("part_agriculture_2023")),
            ("Industrie (%)", r.get("part_industrie_2023")),
            ("Construction (%)", r.get("part_construction_2023")),
            ("Commerce / Transports / Services (%)", r.get("part_commerce_transp_2023")),
            ("Admin. / Santé / Action sociale (%)", r.get("part_admin_sante_2023")),
            ("Établissements 1-9 salariés (%)", r.get("part_etab_1_9_sal_2023")),
            ("Établissements 10+ salariés (%)", r.get("part_etab_10_sal_plus_2023")),
        ]
        _table_2col(doc, rows)
    _section_sep(doc)


def _section_pharmacies(doc, pj_data):
    _heading(doc, "3. Pharmacies et magasins médicaux", 1)
    _heading(doc, "Source : Pages Jaunes", 2)

    total_ph = sum(r.get("nb_pharmacies", 0) for r in pj_data)
    total_mm = sum(r.get("nb_materiel_medical", 0) for r in pj_data)
    _bullet(doc, f"Total pharmacies sur la zone : {total_ph}")
    _bullet(doc, f"Total magasins matériel médical sur la zone : {total_mm}")
    doc.add_paragraph()

    for r in pj_data:
        _heading(doc, r["commune"] + f" ({r['cp']})", 3)
        _bullet(doc, f"Pharmacies : {r.get('nb_pharmacies', 0)}")
        if r.get("noms_pharmacies"):
            for nom in r["noms_pharmacies"]:
                _bullet(doc, nom, level=1)
        _bullet(doc, f"Matériel médical : {r.get('nb_materiel_medical', 0)}")
        if r.get("noms_materiel_medical"):
            for nom in r["noms_materiel_medical"]:
                _bullet(doc, nom, level=1)
    _section_sep(doc)


def _section_maternites(doc, maternites, communes_data):
    _heading(doc, "4. Naissances et maternités", 1)
    total_naissances = sum(
        int(str(r.get("naissances_2022", 0) or 0).replace(" ", "") or 0)
        for r in communes_data
    )
    _bullet(doc, f"Naissances domiciliées 2022 (total zone) : {total_naissances}")
    doc.add_paragraph()
    _heading(doc, "Maternités (source : Journal des Femmes)", 2)
    for m in maternites:
        line = m.get("nom", "")
        if m.get("statut"):
            line += f" — {m['statut']}"
        if m.get("type_niveau"):
            line += f" (Niveau {m['type_niveau']})"
        if m.get("nb_accouchements_an"):
            line += f" — {m['nb_accouchements_an']} accouchements/an"
        if m.get("ville"):
            line += f" — {m['ville']}"
        _bullet(doc, line)
    _section_sep(doc)


def _section_lactariums(doc, lactariums):
    _heading(doc, "5. Lactariums", 1)
    _heading(doc, "Source : Association des Lactariums de France", 2)
    for lac in lactariums:
        line = lac.get("nom", "")
        if lac.get("telephone"):
            line += f" — Tél. {lac['telephone']}"
        _bullet(doc, line)
    _section_sep(doc)


def _section_sages_femmes(doc, sages_femmes):
    _heading(doc, "6. Sages-femmes libérales", 1)
    _heading(doc, f"Source : Ordre national des sages-femmes — {len(sages_femmes)} sage(s)-femme(s) (après dédoublonnage)", 2)
    for sf in sages_femmes:
        parts = [f"{sf.get('nom', '')} {sf.get('prenom', '')}".strip()]
        if sf.get("adresse"):
            parts.append(sf["adresse"])
        if sf.get("telephone"):
            parts.append(f"Tél. {sf['telephone']}")
        if sf.get("email"):
            parts.append(sf["email"])
        _bullet(doc, " | ".join(parts))
    _section_sep(doc)


def _section_pmi(doc, pmi):
    _heading(doc, "7. Protection Maternelle et Infantile (PMI)", 1)
    _heading(doc, "Source : AlloPMI.fr", 2)
    for p in pmi:
        line = p.get("nom", "")
        if p.get("adresse"):
            line += f" — {p['adresse']}"
        if p.get("telephone"):
            line += f" — Tél. {p['telephone']}"
        _bullet(doc, line)
    _section_sep(doc)


def _section_allaitement(doc, dept_nom):
    _heading(doc, "8. Taux d'allaitement (DREES)", 1)
    doc.add_paragraph(
        f"Les données nationales DREES (Enquête nationale périnatale) indiquent :\n"
        f"• Taux d'allaitement à la naissance : entre 57 % et 68 % selon les départements.\n"
        f"• Taux d'allaitement à 10 semaines : entre 27 % et 38 %.\n\n"
        f"Pour les données spécifiques au département {dept_nom}, consulter :\n"
        f"https://drees.solidarites-sante.gouv.fr/publications-documents-de-reference/rapports-etudes/les-naissances-en-2021-et-leurs-caracteristiques"
    )
    doc.add_paragraph(
        f"\nNote de mise à jour : données collectées le {time.strftime('%d/%m/%Y')}. "
        "Millésimes : INSEE RP 2022, DGI 2021, REE 2023."
    )


# ── Point d'entrée ────────────────────────────────────────────────────────────

def generate_word(
    output_path: str,
    zone_nom: str,
    dept_code: str,
    dept_nom: str,
    region: str,
    communes: list[dict],
    communes_data: list[dict],
    pj_data: list[dict],
    maternites: list[dict],
    lactariums: list[dict],
    sages_femmes: list[dict],
    pmi: list[dict],
) -> None:
    """Génère le document Word de synthèse DIP."""
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titre
    title = doc.add_heading(f"SYNTHÈSE DES DONNÉES SOCIO-DÉMOGRAPHIQUES — DIP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Zone : {zone_nom}   |   {dept_nom} ({dept_code})   |   {region}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _section_zone(doc, zone_nom, dept_code, dept_nom, region, communes)
    _section_socio_demo(doc, communes_data)
    _section_pharmacies(doc, pj_data)
    _section_maternites(doc, maternites, communes_data)
    _section_lactariums(doc, lactariums)
    _section_sages_femmes(doc, sages_femmes)
    _section_pmi(doc, pmi)
    _section_allaitement(doc, dept_nom)

    doc.save(output_path)
